from math import floor
from docx.shared import Pt
from docx import Document
from docx.table import Table
from Models.Cliente import Cliente
from Models.Presupuesto import Presupuesto_Lamina
from Services.linker import linker, saver


# for testing purposes
# from test_client import Marcos
# from test_presu import PL


def generar_presupuesto_lamina(cliente: "Cliente", presupuesto: "Presupuesto_Lamina") -> True:

    document_name: str = "06 - Presupuesto Lámina Armada -f.docx"
    document = Document(linker(document_name))
    table = document.tables[0]

    write_client_lamina(cliente, table, presupuesto)
    modelo(document, presupuesto)
    primera_partida(document, presupuesto)
    segunda_partida(document, presupuesto)
    tercera_partida(document, presupuesto)
    totales(document, presupuesto)
    porcentajes(document, presupuesto)

    document.save(saver(document_name))

    return True


def write_client_lamina(cliente: "Cliente", table: "Table", presupuesto: "Presupuesto_Lamina") -> None:

    table.cell(0, 1).paragraphs[0].add_run(f"{cliente.nombre}").bold = True
    table.cell(1, 1).paragraphs[0].add_run(f"{cliente.telefono}").bold = True
    table.cell(1, 3).paragraphs[0].add_run(f"{cliente.nif}").bold = True
    table.cell(2, 1).paragraphs[0].add_run(
        f"{cliente.direccion}, {cliente.localidad}, {cliente.provincia}."
    ).bold = True
    table.cell(3, 1).paragraphs[0].add_run(f"{cliente.correo}").bold = True
    table.cell(4, 1).paragraphs[0].add_run(
        "INGRESAR NUMERO"
    ).bold = True
    table.cell(4, 3).paragraphs[0].add_run(
        f"{presupuesto.fecha.strftime('%x')}"
    ).bold = True

    return


def modelo(document, presupuesto: "Presupuesto_Lamina") -> None:

    document.paragraphs[17].add_run(f" {presupuesto.modelo}").italic = True

    return


def primera_partida(document, presupuesto: "Presupuesto_Lamina") -> None:

    document.paragraphs[7].add_run(f" {presupuesto.metros}m2.").bold = True
    document.paragraphs[7].runs[-1].font.size = Pt(14)

    document.paragraphs[23].add_run(f"      {presupuesto.tarifa}€").bold = True
    document.paragraphs[23].runs[-1].font.size = Pt(14)

    document.paragraphs[25].add_run(
        f"                  {presupuesto.metros}m2"
    ).bold = True
    document.paragraphs[25].runs[-1].font.size = Pt(14)

    document.paragraphs[27].add_run(
        f" {presupuesto.tarifa * presupuesto.metros} €"
    ).bold = True
    document.paragraphs[27].runs[-1].font.size = Pt(14)

    return


def segunda_partida(document, presupuesto: "Presupuesto_Lamina") -> None:

    temp = None
    if presupuesto.impulsion:

        temp = document.paragraphs[37]

        temp.add_run(f"•{chr(9)}Boquillas de impulsión")
        temp.add_run("                ").underline = True
        temp.add_run(
            f"{presupuesto.impulsion_precio} € und. x{presupuesto.impulsion} --------------------------------> {presupuesto.impulsion_precio * presupuesto.impulsion} €"
        )
        temp.style = "List Paragraph"
        temp.runs[0].font.bold = True
        temp.runs[0].font.size = Pt(12)
        temp.runs[1].font.bold = True
        temp.runs[1].font.size = Pt(12)
        temp.runs[2].font.bold = True
        temp.runs[2].font.size = Pt(12)

    if presupuesto.barrefondo:

        temp = document.paragraphs[38]

        temp.add_run(f"•{chr(9)}Boquilla barrefondo")
        temp.add_run("                     ").underline = True
        temp.add_run(
            f"{presupuesto.barrefondo_precio} € und. x{presupuesto.barrefondo} --------------------------------> {presupuesto.barrefondo_precio * presupuesto.barrefondo} €"
        )
        temp.style = "List Paragraph"
        temp.runs[0].font.bold = True
        temp.runs[0].font.size = Pt(12)
        temp.runs[1].font.bold = True
        temp.runs[1].font.size = Pt(12)
        temp.runs[2].font.bold = True
        temp.runs[2].font.size = Pt(12)

    if presupuesto.skimmers:

        temp = document.paragraphs[39]

        temp.add_run(f"•{chr(9)}Skimmers")
        temp.add_run("                                      ").underline = True
        temp.add_run(
            f"{presupuesto.skimmers_precio} € und. x{presupuesto.skimmers} --------------------------------> {presupuesto.skimmers_precio * presupuesto.skimmers} €"
        )
        temp.style = "List Paragraph"
        temp.runs[0].font.bold = True
        temp.runs[0].font.size = Pt(12)
        temp.runs[1].font.bold = True
        temp.runs[1].font.size = Pt(12)
        temp.runs[2].font.bold = True
        temp.runs[2].font.size = Pt(12)

    if presupuesto.sumidero:

        temp = document.paragraphs[40]

        temp.add_run(f"•{chr(9)}Sumidero de Fondo")
        temp.add_run("                    ").underline = True
        temp.add_run(
            f"{presupuesto.sumidero_precio} € und. x{presupuesto.sumidero} --------------------------------> {presupuesto.sumidero_precio * presupuesto.sumidero} €"
        )
        temp.style = "List Paragraph"
        temp.runs[0].font.bold = True
        temp.runs[0].font.size = Pt(12)
        temp.runs[1].font.bold = True
        temp.runs[1].font.size = Pt(12)
        temp.runs[2].font.bold = True
        temp.runs[2].font.size = Pt(12)

    if presupuesto.sumidero_grande:

        temp = document.paragraphs[40]

        temp.add_run(f"•{chr(9)}Sumidero de Fondo")
        temp.add_run("                    ").underline = True
        temp.add_run(
            f"{presupuesto.sumidero_grande_precio} € und. x{presupuesto.sumidero_grande} --------------------------------> {presupuesto.sumidero_grande_precio * presupuesto.sumidero_grande} €"
        )
        temp.style = "List Paragraph"
        temp.runs[0].font.bold = True
        temp.runs[0].font.size = Pt(12)
        temp.runs[1].font.bold = True
        temp.runs[1].font.size = Pt(12)
        temp.runs[2].font.bold = True
        temp.runs[2].font.size = Pt(12)

    if presupuesto.nicho:

        temp = document.paragraphs[41]

        temp.add_run(f"•{chr(9)}Foco + Nicho")
        temp.add_run("                                ").underline = True
        temp.add_run(
            f"{presupuesto.nicho_precio} € und. x{presupuesto.nicho} --------------------------------> {presupuesto.nicho_precio * presupuesto.nicho} €"
        )
        temp.style = "List Paragraph"
        temp.runs[0].font.bold = True
        temp.runs[0].font.size = Pt(12)
        temp.runs[1].font.bold = True
        temp.runs[1].font.size = Pt(12)
        temp.runs[2].font.bold = True
        temp.runs[2].font.size = Pt(12)

    temp = document.paragraphs[44]
    temp.add_run(f" {presupuesto.segunda_partida}€").bold = True
    temp.runs[-1].font.size = Pt(14)

    return


def tercera_partida(document, presupuesto: "Presupuesto_Lamina") -> None:

    if not presupuesto.tercera_partida:
        document.paragraphs[49].add_run("No Procede").bold = True
        return

    temp = document.paragraphs[49]

    temp.add_run(f"{chr(9)}{presupuesto.otros_concepto}")
    temp.add_run("                                ").underline = True
    temp.add_run(
        f"{presupuesto.otros_precio} € und. x{presupuesto.otros} --------------------------------> {presupuesto.otros_precio * presupuesto.otros}€"
    )

    temp.runs[0].font.size = Pt(12)
    temp.runs[0].font.bold = True
    temp.runs[1].font.size = Pt(12)
    temp.runs[1].font.bold = True
    temp.runs[2].font.size = Pt(12)
    temp.runs[2].font.bold = True

    temp = document.paragraphs[51]

    temp.add_run(f" {presupuesto.tercera_partida}€").font.bold = True
    temp.runs[10].font.size = Pt(14)

    return


def totales(document, presupuesto: "Presupuesto_Lamina") -> None:

    temp = document.paragraphs[57]

    temp.add_run(f"{presupuesto.primera_partida}€").bold = True
    temp.runs[-1].font.size = Pt(14)

    temp = document.paragraphs[58]

    temp.add_run(f"{presupuesto.segunda_partida}€").bold = True
    temp.runs[-1].font.size = Pt(14)

    if presupuesto.tercera_partida:

        temp = document.paragraphs[59]

        temp.add_run(f"{presupuesto.tercera_partida}€").bold = True
        temp.runs[-1].font.size = Pt(14)
    else:

        document.paragraphs[59].text = ""

    temp = document.paragraphs[60]

    temp.add_run(
        f"{floor(((presupuesto.total - presupuesto.total/1.21)*100)/100.0)}€"
    ).bold = True
    temp.runs[-1].font.size = Pt(14)

    temp = document.paragraphs[62]

    temp.add_run(f"{presupuesto.total}€").bold = True
    temp.runs[-1].font.size = Pt(18)


def porcentajes(document, presupuesto: "Presupuesto_Lamina") -> None:

    # to do: por alguna razon esto dejo de funcionar bien
    temp = document.paragraphs[74]
    temp.runs[4].text = str(floor((presupuesto.total * 0.8 * 100) / 100.0))
    temp.runs[4].font.size = Pt(14)
    temp.runs[5].text = f"({str(floor(presupuesto.total * .2*100)/100.0)}€)"
    temp.runs[5].font.size = Pt(14)

    return
