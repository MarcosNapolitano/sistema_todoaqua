import datetime
from math import floor
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from Utils.linker import linker, saver


# for testing purposes
# from test_client import Marcos
# from test_presu import PTO


def generar_presupuesto_todo_fugas(cliente: object, presupuesto: object) -> True: 

	document_name: str = "05 - Presupuesto Todo Fugas.docx"
	document: object = Document(linker(document_name))
	date: str = f" {presupuesto.fecha.strftime('%x')}"
	font_color = document.paragraphs[7].runs[0].font.color.rgb

	
	# might want to refactor this
	document.paragraphs[7].add_run(date).font.color.rgb = font_color
	document.paragraphs[7].runs[1].bold = True
	document.paragraphs[8].add_run(f" {cliente.nombre}").font.color.rgb = font_color
	document.paragraphs[8].runs[1].bold = True
	document.paragraphs[9].add_run(f" {cliente.telefono}").font.color.rgb = font_color
	document.paragraphs[9].runs[2].bold = True
	document.paragraphs[10].add_run(f" {cliente.direccion}").font.color.rgb = font_color
	document.paragraphs[10].runs[1].bold = True
	document.paragraphs[11].add_run(f" {cliente.localidad}" +" "*10 + 
								  f"CP: {cliente.cp}" + " "*10 +
								  f"PROVINCIA: {cliente.provincia}").font.color.rgb = font_color
	document.paragraphs[11].runs[1].bold = True

	match presupuesto.tipo:

		# hay que mejorar esto luego
		case "domestica":
			document.paragraphs[17].add_run(" ✔").font.color.rgb = font_color
			document.paragraphs[19].add_run(" X").font.color.rgb = font_color
			document.paragraphs[21].add_run(" X").font.color.rgb = font_color

		case "urbanizacion":
			document.paragraphs[17].add_run(" X").font.color.rgb = font_color
			document.paragraphs[19].add_run(" ✔").font.color.rgb = font_color
			document.paragraphs[21].add_run(" X").font.color.rgb = font_color

		case "comercial":
			document.paragraphs[17].add_run(" X").font.color.rgb = font_color
			document.paragraphs[19].add_run(" X").font.color.rgb = font_color
			document.paragraphs[21].add_run(" ✔").font.color.rgb = font_color

	document.paragraphs[39].add_run(f"{presupuesto.precio}").font.color.rgb = font_color
	document.paragraphs[40].add_run(f"{floor((presupuesto.total - presupuesto.precio) * 100)/100.0}0").font.color.rgb = font_color
	document.paragraphs[41].add_run(f"{presupuesto.total}").font.color.rgb = font_color

	document.paragraphs[41].runs[2].font.size = Pt(16)
	
	document.save(saver(document_name))

	return True