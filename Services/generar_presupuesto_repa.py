import datetime
from docx.shared import Pt
from docx import Document
from docx.table import Table
from Models.Cliente import Cliente
from Models.Presupuesto import Presupuesto_Repa
from Services.linker import linker, saver


# for testing purposes
# from test_client import Marcos
# from test_presu import PR


def generar_presupuesto_repa(cliente: "Cliente", presupuesto: "Presupuesto_Repa") -> True:

    document_name = "08 - Presupuesto Reparación de Piscinas -f.docx"
    document = Document(linker(document_name))
    table = document.tables[0]

    write_client_lamina(cliente, table, presupuesto)
    concepto(document, presupuesto)
    totales(document, presupuesto)

    document.save(saver(document_name))

    return True


def write_client_lamina(cliente: "Cliente", table: "Table", presupuesto: "Presupuesto_Repa") -> None:

    table.cell(0, 1).paragraphs[0].add_run(f"{cliente.nombre}").bold = True
    table.cell(1, 1).paragraphs[0].add_run(f"{cliente.telefono}").bold = True
    table.cell(1, 3).paragraphs[0].add_run(f"{cliente.nif}").bold = True
    table.cell(2, 1).paragraphs[0].add_run(
        f"{cliente.direccion}, {cliente.localidad}, {cliente.provincia}."
    ).bold = True
    table.cell(3, 1).paragraphs[0].add_run(f"{cliente.correo}").bold = True
    table.cell(4, 1).paragraphs[0].add_run(
        "Presupuesto Nº el que sea wachin"
    ).bold = True
    table.cell(4, 3).paragraphs[0].add_run(
        f"{presupuesto.fecha.strftime('%x')}"
    ).bold = True

    return


def concepto(document, presupuesto: "Presupuesto_Repa") -> None:

    index = 10

    # ver forma de hacer dinamico esto porque no calcula bien el espacio

    for i in presupuesto.concepto:

        document.paragraphs[index].add_run(f"• {i[0]}")
        document.paragraphs[index].add_run(" " * 100).underline = True
        document.paragraphs[index].add_run(f"{i[1]}0€")

        document.paragraphs[index].runs[0].font.size = Pt(14)
        document.paragraphs[index].runs[1].font.size = Pt(14)
        document.paragraphs[index].runs[2].font.size = Pt(14)

        index += 1

    return


def totales(document, presupuesto: "Presupuesto_Repa") -> None:

    temp = document.paragraphs[16]

    temp.add_run(f"{presupuesto.precio}€").bold = True
    temp.runs[-1].font.size = Pt(14)

    temp = document.paragraphs[17]

    temp.add_run(f"{presupuesto.total - presupuesto.precio}€").bold = True
    temp.runs[-1].font.size = Pt(14)

    temp = document.paragraphs[18]

    temp.add_run(f"{presupuesto.total}€").bold = True
    temp.runs[-1].font.size = Pt(18)

    return
