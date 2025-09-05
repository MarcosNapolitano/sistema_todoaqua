from math import floor
from docx import Document
from docx.shared import Pt
from Services.linker import linker, saver


# for testing purposes
# from test_client import Marcos
# from test_presu import PF


def generar_presupuesto_fuga_piscinas(cliente: object, presupuesto: object) -> True: 


	document_name: str = "04 - Presupuesto Fugas Piscinas.docx"
	document: object = Document(linker(document_name))
	date: str = f" {presupuesto.fecha.strftime('%x')}"
	font_color = document.paragraphs[4].runs[0].font.color.rgb

	
	# might want to refactor this
	document.paragraphs[4].add_run(date).font.color.rgb = font_color
	document.paragraphs[4].runs[1].bold = True
	document.paragraphs[5].add_run(f" {cliente.nombre}").font.color.rgb = font_color
	document.paragraphs[5].runs[1].bold = True
	document.paragraphs[6].add_run(f" {cliente.telefono}").font.color.rgb = font_color
	document.paragraphs[6].runs[2].bold = True
	document.paragraphs[7].add_run(f" {cliente.direccion}").font.color.rgb = font_color
	document.paragraphs[7].runs[1].bold = True
	document.paragraphs[8].add_run(f" {cliente.localidad}" +" "*10 + 
								  f"CP: {cliente.cp}" + " "*10 +
								  f"PROVINCIA: {cliente.provincia}").font.color.rgb = font_color
	document.paragraphs[8].runs[1].bold = True

	document.paragraphs[18].add_run(f" {presupuesto.ancho}x{presupuesto.largo}m.").font.color.rgb = font_color
	
	if presupuesto.jacuzzi:
		document.paragraphs[20].add_run(" Si").font.color.rgb = font_color
	else:
		document.paragraphs[20].add_run(" No").font.color.rgb = font_color

	if presupuesto.skimmer:
		document.paragraphs[22].add_run(" Skimmer").font.color.rgb = font_color
	else:
		document.paragraphs[20].add_run(" Desbordante").font.color.rgb = font_color


	document.paragraphs[48].add_run(f"{presupuesto.precio}").font.color.rgb = font_color
	document.paragraphs[49].add_run(f"{floor((presupuesto.total - presupuesto.precio) * 100)/100.0}0").font.color.rgb = font_color
	document.paragraphs[50].add_run(f"{presupuesto.total}").font.color.rgb = font_color

	document.paragraphs[50].runs[2].font.size = Pt(16)

	document.save(saver(document_name))

	return True
